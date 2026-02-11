import logging
import os
import pandas as pd
import sys

# Configure logging
logger = logging.getLogger(__name__)

def load_data_high_performance(file_path: str) -> pd.DataFrame:
    """
    Universal high-performance file loader for large datasets (1M+ records).
    
    Industry Best Practice (2026+):
    - Structured Data: Uses the Polars engine (Rust-backed) for 10-50x speedups.
    - Unstructured Data: The 2026+ industry standard is AI-driven extraction (e.g., using
      multimodal LLMs), which is handled via custom parser functions in this framework.
      This loader provides basic, non-AI fallbacks for table extraction from PDF and DOCX.
    
    Supported Formats:
    - Excel (.xlsx, .xls, .xlsb): Uses 'calamine' (Rust-based, memory-mapped).
    - CSV/PSV/TXT: Uses Polars multi-threaded CSV reader.
    - Parquet: Native Arrow/Polars format.
    - Feather/Arrow IPC (.arrow, .feather): Native Arrow/Polars format.
    - Avro (.avro): High-performance binary format.
    - JSON/NDJSON: Optimized for structured and newline-delimited JSON.
    - XML (.xml): Experimental Polars XML reader.
    - PDF (.pdf): Basic table extraction via pdfplumber (slow, not for large volumes).
    - DOCX (.docx): Basic table extraction via python-docx (slow, not for large volumes).
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        import polars as pl
    except ImportError:
        error_msg = "High-performance library 'polars' missing. Please install via: pip install polars"
        logger.error(error_msg)
        raise ImportError(error_msg)

    logger.info(f"Starting high-performance load of '{os.path.basename(file_path)}' using Polars engine...")

    df_pl = None

    # 1. Excel (Rust-based Calamine engine)
    if ext in ['.xlsx', '.xls', '.xlsb', '.xlsm']:
        try:
            import fastexcel # Required for engine='calamine'
        except ImportError:
            raise ImportError("Library 'fastexcel' is required for high-performance Excel parsing.")
        
        # engine='calamine' is the key to 2026+ performance standards
        df_pl = pl.read_excel(source=file_path, engine="calamine")

    # 2. CSV / Text (Multi-threaded SIMD reader)
    elif ext in ['.csv', '.txt', '.psv']:
        # Polars automatically handles CSV parsing much faster than pandas
        try:
            df_pl = pl.read_csv(file_path, ignore_errors=True, try_parse_dates=True)
        except Exception:
            # Fallback for common delimiter issues
            df_pl = pl.read_csv(file_path, separator='\t', ignore_errors=True, try_parse_dates=True)

    # 3. Parquet (Native format for Arrow/Polars)
    elif ext == '.parquet':
        df_pl = pl.read_parquet(file_path)

    # 4. JSON (Structured)
    elif ext in ['.json', '.ndjson', '.jsonl']:
        try:
            # Try Newline Delimited JSON first (common for large data)
            df_pl = pl.read_ndjson(file_path)
        except Exception:
            df_pl = pl.read_json(file_path)

    # 5. High-Performance Binary Formats (Arrow/Feather, Avro)
    elif ext in ['.arrow', '.feather']:
        df_pl = pl.read_ipc(file_path)
    
    elif ext == '.avro':
        df_pl = pl.read_avro(file_path)

    # 6. XML (Semi-structured)
    elif ext == '.xml':
        # Polars' XML reader is powerful but may require specific XPath expressions
        # for complex files. This is a best-effort generic read.
        try:
            df_pl = pl.read_xml(file_path)
        except Exception as e:
            logger.error(f"Polars failed to auto-parse XML. For complex XML, a custom parser with a specific XPath is recommended. Error: {e}")
            raise

    # 7. Unstructured Document Fallbacks (PDF, DOCX)
    elif ext == '.pdf':
        return _load_pdf_tables(file_path)
    
    elif ext == '.docx':
        return _load_docx_tables(file_path)

    else:
        raise ValueError(f"Unsupported file format for high-performance loader: '{ext}' (File: {os.path.basename(file_path)}). Supported formats: Excel, CSV/TXT, Parquet, Arrow/Feather, Avro, JSON, XML, PDF, DOCX.")

    row_count = len(df_pl)
    logger.info(f"Loaded {row_count} rows. Converting to Pandas for pipeline compatibility...")
    
    # Convert to Pandas for downstream compatibility
    # Uses pyarrow for zero-copy conversion if installed
    return df_pl.to_pandas(use_pyarrow_extension_array=True) if row_count > 0 else pd.DataFrame()

def _load_pdf_tables(file_path: str) -> pd.DataFrame:
    """
    Helper to extract tables from PDF.
    Performance Warning: PDF is not a data interchange format. This is slow.
    The 2026+ industry standard is AI-driven extraction for complex documents.
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Library 'pdfplumber' is required for PDF parsing.")

    logger.warning("Parsing PDF with 1M+ records is inefficient. Recommended: Convert to CSV/Parquet upstream.")
    
    all_rows = []
    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)
        logger.info(f"Processing {total_pages} pages of PDF data...")
        
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            for table in tables:
                if not table: continue
                for row in table:
                    # Clean newlines which often break data structure
                    cleaned_row = [cell.replace('\n', ' ').strip() if cell else None for cell in row]
                    all_rows.append(cleaned_row)
            
            if (i + 1) % 50 == 0:
                logger.info(f"Processed {i + 1}/{total_pages} pages...")

    if not all_rows:
        return pd.DataFrame()

    # Assume first row is header
    headers = all_rows[0]
    # Deduplicate headers to prevent Pandas errors
    seen = {}
    deduped_headers = []
    for h in headers:
        h_str = str(h) if h else "Column"
        if h_str in seen:
            seen[h_str] += 1
            deduped_headers.append(f"{h_str}_{seen[h_str]}")
        else:
            seen[h_str] = 0
            deduped_headers.append(h_str)

    return pd.DataFrame(all_rows[1:], columns=deduped_headers)

def _load_docx_tables(file_path: str) -> pd.DataFrame:
    """
    Helper to extract tables from DOCX files.
    Performance Warning: Like PDF, DOCX is not a data interchange format. This is slow.
    The 2026+ industry standard is AI-driven extraction for complex documents.
    """
    try:
        import docx
    except ImportError:
        raise ImportError("Library 'python-docx' is required for DOCX parsing. Install via: pip install python-docx")

    logger.warning("Parsing DOCX with 1M+ records is highly inefficient. Recommended: Convert to a structured format like CSV/Parquet upstream.")
    
    doc = docx.Document(file_path)
    all_rows = []
    
    if not doc.tables:
        logger.warning(f"No tables found in DOCX file: {os.path.basename(file_path)}")
        return pd.DataFrame()
        
    logger.info(f"Processing {len(doc.tables)} tables from DOCX file...")

    # Assume all tables should be concatenated
    for table in doc.tables:
        for row in table.rows:
            # Extract text from each cell in the row
            cleaned_row = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            all_rows.append(cleaned_row)

    if not all_rows:
        return pd.DataFrame()

    # Assume first row of the first table is the header for all concatenated data
    headers = all_rows[0]
    # Deduplicate headers to prevent Pandas errors
    seen = {}
    deduped_headers = []
    for h in headers:
        h_str = str(h) if h else "Column"
        if h_str in seen:
            seen[h_str] += 1
            deduped_headers.append(f"{h_str}_{seen[h_str]}")
        else:
            seen[h_str] = 0
            deduped_headers.append(h_str)

    return pd.DataFrame(all_rows[1:], columns=deduped_headers)

if __name__ == "__main__":
    # CLI Entry point for testing the mechanism
    import time
    logging.basicConfig(level=logging.INFO)
    if len(sys.argv) < 2:
        print("Usage: python fast_data_loader.py <file_path>")
    else:
        start_time = time.time()
        df = load_data_high_performance(sys.argv[1])
        elapsed = time.time() - start_time
        print(f"Success! Loaded {len(df)} rows in {elapsed:.4f} seconds.")
        print(df.head())